from util.GPT import find_sensitive_data, new_file_path
from docx import Document
import streamlit as st
    
def getSensitiveWords(doc):
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)

    highlighted = find_sensitive_data('\n'.join(fullText))
    return highlighted

# Function to remove specified words from the document
def sanitize_docx(uploadedfile):
    import json
    # Open the existing Word document
    doc = Document(uploadedfile)
    words = getSensitiveWords(doc)
    # Remove newline characters and extra spaces
    single_line_json = words.replace('\n', '').replace('  ', '')
    json_NEW = json.loads(single_line_json)
    word_identified = []
    for key in json_NEW:
        if json_NEW[key] != "N/A":
        
            for string in json_NEW[key]:
                word_identified.append(string)
          
                
    st.write(word_identified)
    for paragraph in doc.paragraphs:
        for word in word_identified:
            if word in paragraph.text:
                # Replace the word with an empty string
                paragraph.text = paragraph.text.replace(word, '')
                
    new_filename = new_file_path(uploadedfile.name)
    
  
    # can't use st.download_button here as it doesn't work with docx files
    # so we use a workaround
    with open(new_filename, 'wb') as f:
        doc.save(f)
    with open(new_filename, 'rb') as f:
        bytes = f.read()
        st.download_button(label="Download Redacted File", data=bytes, file_name=new_filename)
    # import os
    # delete the file from the server after download
    # os.remove(new_filename)
    






